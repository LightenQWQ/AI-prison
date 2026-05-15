import os
import glob
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

desktop = os.path.join(os.path.expanduser("~"), "Desktop")
# 尋找原始複製檔案作為基底
matches = glob.glob(os.path.join(desktop, "*只是一個建議*複製*.docx"))

def set_font(run, font_name, size=Pt(12)):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = size

def add_academic_para(doc, text, is_heading=False, level=1):
    if is_heading:
        p = doc.add_heading('', level=level)
        run = p.add_run(text)
        set_font(run, 'DFKai-SB', Pt(16 if level==1 else 14))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level==1 else WD_ALIGN_PARAGRAPH.LEFT
    else:
        p = doc.add_paragraph()
        run = p.add_run("　　" + text) # 學術縮排
        set_font(run, 'PMingLiU', Pt(12))
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

if matches:
    path = matches[0]
    doc = Document(path)
    
    # --- 第四章 ---
    doc.add_page_break()
    add_academic_para(doc, "第四章、創作實踐與系統實作", is_heading=True, level=1)
    add_academic_para(doc, "本章詳述《只是一個建議》從技術原型到完整遊戲系統的開發歷程。本創作旨在建立一個能處理複雜心理敘事且具備美學一致性的 AI 互動環境。由於本作核心在於結合大型語言模型（LLM）的即時推理能力與漫畫式視覺生成方式，因此系統架構必須在維持高穩定性的同時，同步支援敘事解析、風格化濾鏡轉譯、以及跨平台的數據流轉。")
    
    add_academic_para(doc, "一、基於 FastAPI 之高併發後端架構與通訊邏輯", is_heading=True, level=2)
    add_academic_para(doc, "本研究採用 FastAPI 作為核心開發環境，建構了一個具備高度擴充性的「雙引擎通訊中樞」。相對於傳統的單機運算，FastAPI 網頁框架能有效整合雲端 AI 模型（如 Gemini 與 Vertex AI）的異步調用，使敘事生成、情緒權重計算與影像產出能並行運作。此架構不僅支援了玩家的自然語言輸入（Open Text Input），更透過自定義的數據模型（Pydantic Models），精確控制了角色的情緒數值（信任、恐懼、懷疑）在不同回合間的連貫性。這使得少年的回應不再是孤立的字串，而是具備歷史記憶與動態心理曲線的數位存在。")
    add_academic_para(doc, "在前端整合層面，本系統利用非同步請求（Async/Await）機制，確保玩家在輸入指令後，系統能即時回饋文字敘事，同時在背景進行高品質的圖像渲染。透過統一的 Google GenAI SDK 介面，本研究實作了「內容與視覺分離」的傳輸策略：文字大腦負責處理語氣中的微小情緒波動，而影像引擎則專注於美學細節的堆疊。這種架構讓開發者能靈活調整模型版本（如從 Imagen 3.0 升級至 4.0），而無需改動核心遊戲邏輯，為長期的創作實驗提供了極佳的彈性與穩定性。")

    add_academic_para(doc, "二、雙引擎整合實踐：Gemini 語意解析與 Vertex AI 視覺轉譯", is_heading=True, level=2)
    add_academic_para(doc, "語言生成流程採用 Gemini 1.5 Flash 作為敘事中樞。玩家輸入的文字會先由系統指令（System Instructions）進行過濾與引導，模型會偵測其中的威脅度、安撫強度或邏輯深度，並將其轉化為少年的心理變量。本研究特別強調「非對稱性互動」，當情緒模型偵測到玩家語氣過於強勢時，會觸發少年的「逆反心理」，產生拒絕執行或憤怒回應的敘事分支。這種處理方式讓角色不再只是單純的指令接收器，而是具有自主意志（Agency）的生命體，其回應語句會隨著信任程度展現出從「防備、敵視」到「依賴、合作」的細微轉變。")
    add_academic_para(doc, "影像生成流程則是以「黑色電影表現主義」為美學框架。系統在接收到邏輯引擎產出的場景指令後，會自動套用 HAMP（Hardboiled Artistic Metaphor Protocol）協定進行視覺轉譯。針對生成式 AI 常見的安全性過濾問題，HAMP 協定透過「藝術隱喻」將物理性的衝突轉化為視覺符號。例如，當劇本出現劇烈情緒波動時，系統會自動在生圖提示詞（Prompt）中加入「不穩定的構圖」與「高反差的陰影欄杆（Shadow Bars）」，利用光影的壓迫感來外化角色的內在焦慮。Vertex AI 產出的 Imagen 3.0 影像，能精確呈現出帶有網點（Screentone）質感的黑白畫面，使每一幀分鏡都能與當下的敘事張力高度同步。")

    add_academic_para(doc, "三、互動式空間優先控制與視覺一致性協定", is_heading=True, level=2)
    add_academic_para(doc, "本研究的技術架構中最具原創性的部分，在於建立了「空間優先（Space-First）」的視覺穩定協定。在即時生圖的環境下，背景的漂移常會導致玩家失去場景定位感，為此，系統強制規定 AI 在生成影像時，必須先建構「冷冽的金屬機械空間、佈滿管線的天花板、具壓迫感的鋼鐵牆面」等環境基調。這種由外而內的生成順序，確保了即使少年的動作在每一回合有所不同，其身處的「密閉空間」依然具備強烈的一致性，進而強化了「困境」這一敘事核心。")
    add_academic_para(doc, "最終，這套架構形成了一個動態的閉環系統：玩家輸入影響 AI 心理，AI 心理驅動文字回應與視覺分鏡，而充滿心理張力的影像與聲音再反過來影響玩家對角色的理解。透過這種跨層次整合方式，遊戲形成一個動態敘事的循環，使黑白心理懸疑漫畫風得以在互動敘事中實現，使玩家在對話中逐步感受角色的情緒波動，並與 AI 共同構築獨特的故事文本。")

    # --- 第五章 ---
    doc.add_page_break()
    add_academic_para(doc, "第五章、研究發現與結論", is_heading=True, level=1)
    
    add_academic_para(doc, "一、敘事能動性之變遷：指引機制對玩家沉浸感之影響", is_heading=True, level=2)
    add_academic_para(doc, "本研究發現，當遊戲機制從「絕對操控」轉向「間接指引」時，玩家的能動性（Agency）並未消失，而是轉化為一種更深層的情感參與。在測試過程中，多數玩家表示，當 AI 角色表現出「因恐懼而拒絕建議」或「因懷疑而反問玩家」時，他們會產生一種與真實生命交流的錯覺。這種基於 Erikson 青少年心理學理論所設計的「情緒不確定性」，成功地將玩家從單純的「解謎者」提升為「引導者」。玩家必須學會調整自己的語氣與邏輯，這種「理解他人心理」的過程，成為了比解謎本身更具吸引力的遊戲核心。")
    add_academic_para(doc, "此外，生成式 AI 帶來的「不可預期性」徹底打破了傳統解謎遊戲的固定劇本。每一回合的回應都是即時生成的，這意味著敘事曲線具備了無限的變異性。研究發現，這種「非線性」的對話流，能激發玩家嘗試各種極端的建議（如恐嚇或極度安撫）以觀察角色的反應，進而產生了一種「敘事共創（Co-creation）」的體驗。玩家不再是被動地消耗內容，而是透過與 AI 的博弈，共同決定了這個 Noir 故事的基調與走向。")

    add_academic_para(doc, "二、視覺修辭與風格穩定性：HAMP 協定與空間美學之驗證", is_heading=True, level=2)
    add_academic_para(doc, "在視覺表現層面，本研究證實了「黑白心理懸疑風格」與生成式 AI 的高度適配性。透過 HAMP 協定的實作，系統成功地將模型受限的物理描述轉化為富有張力的視覺修辭。研究發現，當 AI 使用「深邃的陰影塊」與「銳利的線條抖動」來表現角色的精神壓力時，其敘事強度遠高於傳統的具象化描繪。這種「去色、強化對比、引入符號性」的手法，有效地彌補了現階段生成模型在細部比例上的瑕疵，證明了在 AI 創作中，「風格化的約束」往往比「寫實的追求」更能產生藝術共鳴。")
    add_academic_para(doc, "另一方面，空間優先（Space-First）邏輯的成功實踐，解決了互動生圖領域長期存在的「場景漂移」問題。測試結果顯示，當場景中的工業機械元素（如巨大的齒輪、冷冽的牆面）保持穩定時，玩家對於少年的動作變化具有更高的包容度。這種「環境不動、情緒動」的視覺策略，不僅符合黑色電影（Film Noir）中場景作為心理外化的美學標準，更在技術層面上為即時生成遊戲提供了一套可行的「視覺錨點（Visual Anchor）」解決方案。")

    add_academic_para(doc, "三、結論與未來展望", is_heading=True, level=2)
    add_academic_para(doc, "總結而言，本研究透過《只是一個建議》的創作，成功建立了一套結合 LLM 邏輯推理與高階視覺生成的互動遊戲框架。我們不僅驗證了 AI 作為「具備自主意志之角色」在敘事上的潛力，更提出了一套能克服 API 安全限制與視覺穩定性問題的技術協定。本研究的主要貢獻在於，它示範了創作者如何透過「提示工程」與「美學架構」的介入，在 AI 生成的不確定性中引導出高品質的藝術產出，實現了真正意義上的「人機協作敘事」。")
    add_academic_para(doc, "展望未來，生成式 AI 在遊戲創作中的應用將不僅止於文字與影像。下一階段的研究可聚焦於「多模態感官的即時耦合」，例如開發能根據 AI 角色情緒即時生成的「心理張力音效系統」。隨著計算能力的提升與模型本地化的發展，未來的數位創作將可能實現「完全去劇本化」的動態世界，讓每一段互動都能成為深刻的心理冒險。")

    # --- 參考文獻 ---
    doc.add_page_break()
    add_academic_para(doc, "參考文獻", is_heading=True, level=1)
    add_academic_para(doc, "Erikson, E. H. (1968). Identity: Youth and Crisis. New York: Norton.")
    add_academic_para(doc, "McCloud, S. (1993). Understanding Comics: The Invisible Art. HarperPerennial.")
    add_academic_para(doc, "Riedl, M. O., & Bulitko, V. (2013). Interactive Narrative: An Intelligent Systems Approach. AI Magazine.")
    add_academic_para(doc, "Turner, V. (1969). The Ritual Process: Structure and Anti-Structure. Chicago: Aldine Publishing Co.")

    final_path = os.path.join(desktop, "《只是一個建議》完整論文_深度改寫版.docx")
    doc.save(final_path)
    print(f"SUCCESS: {final_path}")
else:
    print("ERROR: Could not find original template.")
