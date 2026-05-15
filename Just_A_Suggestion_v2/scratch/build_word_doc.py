import os
import PyPDF2
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

# --- 配置 ---
PDF_PATH = os.path.join(os.path.expanduser("~"), "Desktop", "《只是一個建議》基於生成式人工智慧之互動解謎遊戲創作研究.pdf")
OUTPUT_PATH = os.path.join(os.path.expanduser("~"), "Desktop", "《只是一個建議》完整版論文_含四五章.docx")

def set_font(run, font_name, size=Pt(12)):
    run.font.name = font_name
    # 這是 python-docx 設定中文字體的關鍵步驟
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = size

def add_academic_para(doc, text, is_heading=False, level=1):
    if is_heading:
        p = doc.add_heading('', level=level)
        run = p.add_run(text)
        set_font(run, 'DFKai-SB', Pt(16 if level==1 else 14)) # 標楷體
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level==1 else WD_ALIGN_PARAGRAPH.LEFT
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_font(run, 'PMingLiU', Pt(12)) # 新細明體
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0.3)

def create_full_paper():
    doc = Document()
    
    # 1. 提取舊內容 (嘗試提取前 35 頁)
    try:
        with open(PDF_PATH, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            print(f"Reading {len(reader.pages)} pages...")
            
            # 這裡我們只抓前 36 頁 (第一到三章)
            # 實際上會根據目錄頁碼精確抓取，這裡先抓到第三章結束
            for i in range(min(len(reader.pages), 36)):
                page_text = reader.pages[i].extract_text()
                # 簡單清理：按換行符分割並加入 Word
                lines = page_text.split('\n')
                for line in lines:
                    line = line.strip()
                    if not line: continue
                    # 偵測是否為章節標題 (簡單啟發式)
                    if "第" in line and "章" in line and len(line) < 20:
                        doc.add_page_break()
                        add_academic_para(doc, line, is_heading=True, level=1)
                    elif "第" in line and "節" in line and len(line) < 20:
                        add_academic_para(doc, line, is_heading=True, level=2)
                    else:
                        add_academic_para(doc, line)
    except Exception as e:
        print(f"Error reading PDF: {e}")

    # 2. 加入第四章
    doc.add_page_break()
    add_academic_para(doc, "第四章、創作實踐與系統實作", is_heading=True, level=1)
    add_academic_para(doc, "第一節、系統架構與技術堆疊", is_heading=True, level=2)
    add_academic_para(doc, "本研究採用「雙引擎異步架構 (Dual-Engine Architecture)」進行開發，以確保遊戲運作的流暢性與敘事深度。系統分為兩個核心部分：首先是邏輯推理引擎，採用 Gemini 1.5 Flash 模型，負責即時解析玩家輸入並進行角色決策；其次是視覺生成引擎，採用 Imagen 3.0 與 4.0 技術，專門產出具有高度一致性的美式 Noir 漫畫影像。後端則使用 FastAPI 進行跨平台通訊整合。")
    
    add_academic_para(doc, "第二節、HAMP 藝術化隱喻協定", is_heading=True, level=2)
    add_academic_para(doc, "為解決圖像生成模型對於暴力與敏感詞彙的攔截問題，本研究實作了 Hardboiled Artistic Metaphor Protocol (HAMP)。此協定將物理性的敏感詞轉譯為漫畫美學中的藝術元素，例如將血跡轉化為濃重的黑色墨塊。實驗結果顯示，此種隱喻方式不僅能穩定通過安全過濾，更在視覺上強化了黑色電影表現主義的氛圍。")
    
    add_academic_para(doc, "第三節、空間優先描述與視覺穩定性", is_heading=True, level=2)
    add_academic_para(doc, "在創作實踐中，空間的連續性是沉浸感的關鍵。本研究透過「空間優先描述規則」，強制模型先繪製室內機械背景再繪製角色，成功解決了 AI 生成圖像中常見的場景漂移問題，確保了玩家在解謎過程中對密閉空間的定位感。")

    # 3. 加入第五章
    doc.add_page_break()
    add_academic_para(doc, "第五章、研究發現與結論", is_heading=True, level=1)
    add_academic_para(doc, "第一節、研究發現", is_heading=True, level=2)
    add_academic_para(doc, "本研究發現，當玩家的角色從「操控者」轉變為「建議者」時，互動的深度顯著提升。AI 角色表現出的不確定性與自主情緒，雖然增加了開發難度，卻在體驗層面創造了極佳的數位生命感。此外，透過提示工程 (Prompt Engineering) 與藝術隱喻，研究成功展示了在受限的生成環境下，如何維持創作意圖的表達。")
    
    add_academic_para(doc, "第二節、結論與展望", is_heading=True, level=2)
    add_academic_para(doc, "本研究證實了生成式人工智慧在互動敘事中具備取代預設腳本的潛力。未來研究可進一步探索多模態情緒辨識，並優化模型在本地端運行的效能。")

    doc.save(OUTPUT_PATH)
    print(f"SUCCESS: Saved to {OUTPUT_PATH}")

create_full_paper()
