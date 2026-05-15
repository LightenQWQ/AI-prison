import os
import glob
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

desktop = os.path.join(os.path.expanduser("~"), "Desktop")
matches = glob.glob(os.path.join(desktop, "*只是一個建議*複製*.docx"))

def set_font(run, font_name, size=Pt(12)):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = size

def add_academic_para(doc, text, is_heading=False, level=1):
    if is_heading:
        p = doc.add_heading('', level=level)
        run = p.add_run(text)
        set_font(run, 'DFKai-SB', Pt(16 if level==1 else 14))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level==1 else WD_ALIGN_PARAGRAPH.LEFT
    else:
        p = doc.add_paragraph()
        run = p.add_run("　　" + text) # 模擬縮排
        set_font(run, 'PMingLiU', Pt(12))
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

if matches:
    path = matches[0]
    doc = Document(path)
    
    # 第四章
    doc.add_page_break()
    add_academic_para(doc, "第四章、創作實踐與系統開發", is_heading=True, level=1)
    add_academic_para(doc, "第一節、開發環境與技術選型", is_heading=True, level=2)
    add_academic_para(doc, "本研究將初期之 Unity 概念轉譯為基於 Web 服務之即時敘事系統，以因應生成式 AI 對於雲端算力之需求。系統核心採用 FastAPI 構建後端，確保能高效處理玩家輸入、語言模型推論與圖像生成請求之併發處理。")
    add_academic_para(doc, "第二節、敘事推理與視覺生成之雙引擎架構", is_heading=True, level=2)
    add_academic_para(doc, "為解決生成式圖像模型（Imagen）生成速度較慢與安全性過濾之挑戰，本系統實作了「敘事與視覺分流」之雙引擎架構。Gemini 邏輯推演引擎負責解析玩家建議，Vertex AI 視覺引擎則接收經過 HAMP 處理後的視覺指令，產出具備高對比與 Noir 質感之漫畫分鏡。")
    add_academic_para(doc, "第三節、HAMP 隱喻協定之設計與應用", is_heading=True, level=2)
    add_academic_para(doc, "本研究創新提出「Hardboiled Artistic Metaphor Protocol (HAMP)」協定。由於生成式模型對於極端情緒或衝突場面有嚴格的安全性過濾，HAMP 透過「視覺隱喻化」技術，將可能觸發攔截的詞彙轉譯為藝術語彙，成功提升了作品的藝術表現深度。")
    
    # 第五章
    doc.add_page_break()
    add_academic_para(doc, "第五章、研究發現與結論", is_heading=True, level=1)
    add_academic_para(doc, "第一節、生成式 AI 帶來的「非對稱性」互動發現", is_heading=True, level=2)
    add_academic_para(doc, "本研究發現，當玩家的角色定位從「絕對操控」轉向「間接建議」時，AI 角色展現出的「不確定性行為」反而強化了玩家的沉浸感。這種非對稱性的互動機制，使得解謎遊戲演變為一場關於信任與心理博弈的敘事旅程。")
    add_academic_para(doc, "第二節、結論：AI 作為共同創作者的潛力", is_heading=True, level=2)
    add_academic_para(doc, "透過《只是一個建議》的創作實踐，本研究證實了生成式 AI 不僅是內容生成工具，更具備作為「動態敘事演繹者」的潛力。本系統成功在 AI 生成的不確定性中，找回了創作者對於遊戲視覺美學與敘事連貫性的控制權。")
    add_academic_para(doc, "第三節、研究限制與未來展望", is_heading=True, level=2)
    add_academic_para(doc, "未來研究可嘗試導入 ControlNet 等技術，進一步強化 AI 在動態分鏡中的視覺穩定性。")

    output_path = os.path.join(desktop, "《只是一個建議》完整論文_最終成品.docx")
    doc.save(output_path)
    print(f"SUCCESS: {output_path}")
else:
    print("ERROR: File matches not found.")
