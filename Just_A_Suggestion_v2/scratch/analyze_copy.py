import os
from docx import Document

target_name = "《只是一個建議》基於生成式人工智慧之互動解謎遊戲創作研究 - 複製.docx"
path = os.path.join(os.path.expanduser("~"), "Desktop", target_name)

def read_and_analyze():
    if not os.path.exists(path):
        print("ERROR: File not found.")
        return
    
    doc = Document(path)
    print(f"File loaded. Total paragraphs: {len(doc.paragraphs)}")
    
    # 抓取前 20 個段落來分析風格
    for i in range(min(len(doc.paragraphs), 20)):
        p = doc.paragraphs[i]
        if p.text.strip():
            print(f"P{i}: {p.text[:100]}...")

read_and_analyze()
